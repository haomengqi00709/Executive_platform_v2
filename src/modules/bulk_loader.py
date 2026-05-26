"""
bulk_loader — extract CRM contacts or Projects from user-uploaded files.

Pipeline:
  raw bytes (.csv/.xlsx/.docx/.pdf) → textual representation → Gemini prompt → list[dict]

PDFs go through Gemini's file upload (native PDF handling); the other formats
are parsed locally to text/CSV first so the prompt is cheap and deterministic.
"""
import csv
import io
import json
import re
import tempfile
import time
from pathlib import Path

from src.ai import AIClient


# ── Schema field lists ─────────────────────────────────────

_CRM_FIELDS = [
    "email", "name", "company", "role", "phone", "linkedin",
    "status", "priority", "summary", "notes",
]

_PROJECT_FIELDS = [
    "name", "summary", "status", "category", "participants",
    "key_topics", "next_action", "deadline",
]

_CRM_STATUSES   = {"client", "prospect", "partner", "vendor", "other"}
_PROJECT_STATUS = {"ongoing", "needs_attention", "paused", "early_stage", "completed"}


# ── File → text conversion ────────────────────────────────

def _ext(filename: str) -> str:
    return Path(filename).suffix.lower()


def _read_csv(file_bytes: bytes) -> str:
    """Return the CSV content as-is (text)."""
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def _read_xlsx(file_bytes: bytes) -> str:
    """Flatten Excel workbook into CSV-like text (one sheet per block)."""
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
    blocks: list[str] = []
    for sheet in wb.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            if any(c is not None and str(c).strip() for c in row):
                rows.append(",".join("" if c is None else str(c).replace(",", " ") for c in row))
        if rows:
            blocks.append(f"### Sheet: {sheet.title}\n" + "\n".join(rows))
    return "\n\n".join(blocks)


def _read_docx(file_bytes: bytes) -> str:
    """Extract all paragraphs + tables from a .docx file."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    lines: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _read_text(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


def _to_text(file_bytes: bytes, filename: str) -> tuple[str, str]:
    """Convert any supported file to a text representation.
    Returns (text, kind) where kind ∈ csv|xlsx|docx|pdf|txt.
    For PDFs, text is empty — caller should use the Gemini file upload path instead.
    """
    ext = _ext(filename)
    if ext == ".csv":
        return _read_csv(file_bytes), "csv"
    if ext in (".xlsx", ".xls"):
        return _read_xlsx(file_bytes), "xlsx"
    if ext == ".docx":
        return _read_docx(file_bytes), "docx"
    if ext in (".txt", ".md"):
        return _read_text(file_bytes), "txt"
    if ext == ".pdf":
        return "", "pdf"
    raise ValueError(f"Unsupported file extension: {ext}")


# ── AI extraction ─────────────────────────────────────────

def _crm_prompt(content: str) -> str:
    return f"""You are extracting CRM contacts from a user-uploaded file.

Input file content:
---
{content[:30000]}
---

Extract every distinct person/contact. Return a JSON array. Each element must
have these keys (use empty string when unknown, never invent data):

- email     — REQUIRED, valid email address. Skip rows without one.
- name      — Full name. If "Smith, John" reorder to "John Smith".
- company   — Company name.
- role      — Job title.
- phone     — Phone number as written.
- linkedin  — LinkedIn URL.
- status    — Best guess from context: {sorted(_CRM_STATUSES)}. Default "other".
- priority  — "high"|"medium"|"low". Default "medium".
- summary   — One sentence relationship/topic context if mentioned, else "".
- notes     — Any extra remarks not covered above.

Rules:
- Map column headers intelligently ("Email"/"E-mail"/"邮箱" → email, etc.).
- Multilingual headers OK (Chinese, Spanish, etc. — treat them like English equivalents).
- Skip rows that are clearly headers, totals, or notes.
- Skip rows missing both email AND name.
- Output ONLY the JSON array, no markdown fences, no commentary.
"""


def _project_prompt(content: str) -> str:
    return f"""You are extracting business projects/initiatives from a user-uploaded file.

Input file content:
---
{content[:30000]}
---

Extract every distinct project, initiative, deal, or engagement. Return a JSON
array. Each element must have these keys (empty string/[] when unknown):

- name         — REQUIRED, the project title.
- summary      — 1-2 sentence description.
- status       — Best guess from {sorted(_PROJECT_STATUS)}. Default "ongoing".
- category     — "client_deal"|"internal"|"vendor"|"partnership"|"other". Default "other".
- participants — Array of email addresses if any are mentioned.
- key_topics   — Array of short keyword strings (3-8 tokens each).
- next_action  — If a next step is mentioned, capture it in one sentence.
- deadline     — Date string if mentioned, else "".

Rules:
- Skip rows that aren't real projects (headers, blank rows, summaries).
- Map free-form descriptions intelligently — a Word doc paragraph about "Acme
  migration" should still become one project entry.
- Output ONLY the JSON array, no markdown fences, no commentary.
"""


def _parse_json_array(raw: str) -> list[dict]:
    raw = (raw or "").strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
    # Sometimes AI wraps in {"contacts": [...]} — find the array
    try:
        data = json.loads(raw)
    except Exception:
        match = re.search(r"\[.*\]", raw, re.DOTALL)
        if not match:
            return []
        try:
            data = json.loads(match.group())
        except Exception:
            return []
    if isinstance(data, dict):
        # Try common wrapper keys
        for k in ("contacts", "projects", "items", "data"):
            if isinstance(data.get(k), list):
                data = data[k]
                break
        else:
            return []
    if not isinstance(data, list):
        return []
    return [d for d in data if isinstance(d, dict)]


def _run_ai_on_text(ai: AIClient, content: str, prompt_builder) -> list[dict]:
    if not content.strip():
        return []
    prompt = prompt_builder(content)
    raw = ai.generate(prompt)
    return _parse_json_array(raw)


def _run_ai_on_pdf(ai: AIClient, file_bytes: bytes, prompt_builder) -> list[dict]:
    """For PDFs, upload to Gemini Files API and ask the model to read it directly."""
    from google.genai import types as genai_types

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(file_bytes)
        tmp = f.name

    try:
        uploaded = ai.client.files.upload(file=tmp)
        while uploaded.state.name == "PROCESSING":
            time.sleep(2)
            uploaded = ai.client.files.get(name=uploaded.name)

        # The prompt builder expects content text — but for PDFs we pass it as a file.
        # So we feed a placeholder content; the AI reads the file directly.
        prompt = prompt_builder("(see the attached PDF)")

        response = ai.client.models.generate_content(
            model=ai.model,
            contents=[
                genai_types.Part.from_uri(file_uri=uploaded.uri, mime_type="application/pdf"),
                prompt,
            ],
        )
        return _parse_json_array(response.text or "")
    finally:
        Path(tmp).unlink(missing_ok=True)


# ── Normalisation (defensive) ─────────────────────────────

def _normalise_crm(item: dict) -> dict | None:
    email = (item.get("email") or "").strip().lower()
    if not email or "@" not in email:
        return None
    status = (item.get("status") or "other").strip().lower()
    if status not in _CRM_STATUSES:
        status = "other"
    priority = (item.get("priority") or "medium").strip().lower()
    if priority not in ("high", "medium", "low"):
        priority = "medium"
    return {
        "email":    email,
        "name":     str(item.get("name") or "").strip()[:100],
        "company":  str(item.get("company") or "").strip()[:100],
        "role":     str(item.get("role") or "").strip()[:100],
        "phone":    str(item.get("phone") or "").strip()[:50],
        "linkedin": str(item.get("linkedin") or "").strip()[:200],
        "status":   status,
        "priority": priority,
        "summary":  str(item.get("summary") or "").strip()[:500],
        "notes":    str(item.get("notes") or "").strip()[:500],
    }


def _normalise_project(item: dict) -> dict | None:
    name = (item.get("name") or "").strip()
    if not name:
        return None
    status = (item.get("status") or "ongoing").strip().lower()
    if status not in _PROJECT_STATUS:
        status = "ongoing"
    category = (item.get("category") or "other").strip().lower()
    if category not in ("client_deal", "internal", "vendor", "partnership", "other"):
        category = "other"
    participants = item.get("participants") or []
    if not isinstance(participants, list):
        participants = []
    participants = [str(p).strip().lower() for p in participants if str(p).strip()]
    key_topics = item.get("key_topics") or []
    if not isinstance(key_topics, list):
        key_topics = []
    key_topics = [str(t).strip()[:50] for t in key_topics if str(t).strip()]
    return {
        "name":         name[:120],
        "summary":      str(item.get("summary") or "").strip()[:500],
        "status":       status,
        "category":     category,
        "participants": participants[:30],
        "key_topics":   key_topics[:15],
        "next_action":  str(item.get("next_action") or "").strip()[:300],
        "deadline":     str(item.get("deadline") or "").strip()[:30],
    }


# ── Public API ────────────────────────────────────────────

def parse_crm_file(file_bytes: bytes, filename: str, ai: AIClient) -> list[dict]:
    text, kind = _to_text(file_bytes, filename)
    if kind == "pdf":
        raw = _run_ai_on_pdf(ai, file_bytes, _crm_prompt)
    else:
        raw = _run_ai_on_text(ai, text, _crm_prompt)
    out: list[dict] = []
    seen_emails: set = set()
    for r in raw:
        n = _normalise_crm(r)
        if not n:
            continue
        if n["email"] in seen_emails:
            continue
        seen_emails.add(n["email"])
        out.append(n)
    return out


def parse_project_file(file_bytes: bytes, filename: str, ai: AIClient) -> list[dict]:
    text, kind = _to_text(file_bytes, filename)
    if kind == "pdf":
        raw = _run_ai_on_pdf(ai, file_bytes, _project_prompt)
    else:
        raw = _run_ai_on_text(ai, text, _project_prompt)
    out: list[dict] = []
    seen_names: set = set()
    for r in raw:
        n = _normalise_project(r)
        if not n:
            continue
        key = n["name"].lower()
        if key in seen_names:
            continue
        seen_names.add(key)
        out.append(n)
    return out
